"""Build the illustrated per-role operations guide as a Word document.

    pip install python-docx
    python scripts/render_diagrams.py --out ./ie-diagrams
    python scripts/build_role_guide_docx.py \
        --shots ./ie-screenshots --diagrams ./ie-diagrams --out ./Operations-Guide.docx

One chapter per human seat, one section per activity, illustrated with two
kinds of figure:

  diagrams    rendered from docs/diagrams.md, so the pictures in this document
              and the pictures in the repository can never drift apart
  screenshots captured from a live deployment

Screenshots show account-specific values (account id, bucket names, console
URLs). Capture your own from your own deployment rather than reusing someone
else's - SHOT_FILES lists exactly what the document expects, and any missing
file renders as a visible placeholder instead of failing the build.

The prose mirrors docs/guides/. If a guide changes, change this in the same
PR.
"""
from __future__ import annotations

import argparse
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

# --- what the document expects ------------------------------------------------
SHOT_FILES = [
    "01-cloudformation-stacks.jpg",    # CloudFormation stacks, filtered to the project
    "02-lakehouse-tiers.jpg",          # lakehouse bucket root: the medallion tier folders
    "05-stewardship-gate-events.jpg",  # stewardship/gate-events/ in S3
    "07-glue-lineage-properties.jpg",  # Glue table properties: the ie.* lineage keys
    "09-stepfunctions-machine.jpg",    # state machine page + execution history
    "10-stepfunctions-definition.jpg", # definition tab: JSON beside the graph
    "12-execution-graph.png",          # a succeeded execution's graph view
    "13-cloudtrail-trail.jpg",         # trail: validation, multi-region, data events
    "15-consultant-console.jpg",       # consultant home
    "16-engineer-dashboard.jpg",       # /engineer dashboard
    "17-guardrails.jpg",               # /engineer/guardrails
    "18-fde-diff-stale-case.jpg",      # workbench terminal: the agent's diff
    "19-fde-eval-recall.jpg",          # workbench terminal: eval recall against Bedrock
    "20-fde-claude-code-session.jpg",  # Claude Code session: analysis + boundary probe
]

DIAGRAMS = {
    "system": "01-the-whole-system-on-one-page.png",
    "run": "02-a-report-run-end-to-end.png",
    "admission": "03-data-admission-through-the-medallion-tiers.png",
    "deploy": "04-the-deploy-channel.png",
    "release": "05-pipeline-release-candidate-blessed.png",
    "criticality": "06-the-access-model-at-a-glance.png",
    "audit": "07-the-audit-surface.png",
    "analyst": "08-the-analyst-s-zero-copy-path.png",
}

TEAL = RGBColor(0x0D, 0x94, 0x88)
GREY = RGBColor(0x5B, 0x67, 0x7D)
AMBER = RGBColor(0xB4, 0x53, 0x09)

ap = argparse.ArgumentParser()
ap.add_argument("--shots", default="./ie-screenshots")
ap.add_argument("--diagrams", default="./ie-diagrams")
ap.add_argument("--out", default="./Intelligence-Engine-Operations-Guide.docx")
args = ap.parse_args()

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)


# --- helpers ------------------------------------------------------------------
def title(text, size=30, color=TEAL, center=True, space_after=6):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.color.rgb = color
    return p


def h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = TEAL
    return h


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


def para(text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    return p


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def steps(items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    return p


def caption(text):
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(text)
    r.font.size = Pt(8.5)
    r.font.italic = True
    r.font.color.rgb = GREY


def _figure(path, text, width):
    if not os.path.exists(path):
        para(f"[figure missing: {os.path.basename(path)}]", italic=True, color=AMBER)
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(text)


def shot(fname, text, width=6.4):
    _figure(os.path.join(args.shots, fname), text, width)


def diagram(key, text, width=6.0):
    _figure(os.path.join(args.diagrams, DIAGRAMS[key]), text, width)


def table(rows, widths=None):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            r = cells[j].paragraphs[0].add_run(val)
            r.font.size = Pt(9.5)
            r.font.bold = (i == 0)
    if widths:
        for row in t.rows:
            for j, wdt in enumerate(widths):
                row.cells[j].width = Inches(wdt)
    doc.add_paragraph()
    return t


# --- cover --------------------------------------------------------------------
for _ in range(5):
    doc.add_paragraph()
title("Intelligence Engine")
title("Role-by-Role Operations Guide", size=17, color=GREY, space_after=18)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Six human seats, their activities, and the evidence that each one works.\n"
              "Every screenshot is from the live deployment; every diagram is rendered\n"
              "from the same source the repository publishes.")
r.font.size = Pt(11)
r.font.color.rgb = GREY
doc.add_page_break()

# --- orientation --------------------------------------------------------------
h1("Orientation")
para("The Intelligence Engine is a pre-engagement intelligence substrate for "
     "management consultants. Synthetic workforce data flows through a governed "
     "medallion data plane; an episodic LLM pipeline turns it into client "
     "briefings; a hosted console puts a human checkpoint in front of anything "
     "that ships. It runs on a single AWS account, deployed entirely from "
     "CloudFormation templates in a public repository.")
diagram("system", "Figure 1. The whole system: people, product, governed data plane, and the "
                  "governance surfaces that watch all of it.")
para("Two properties are worth stating before anything else, because every "
     "chapter below depends on them.")
bullets([
    "Everything that exists is in a template. Fifteen stacks own every bucket, "
    "role, state machine and trail in the account; a drift detector "
    "(scripts/iac_coverage.py) fails if anything exists outside them.",
    "Read and write are never the same grant. Each medallion tier has separate "
    "read and write policies held by different roles, so no identity can "
    "rewrite the tier it consumes.",
])
shot("01-cloudformation-stacks.jpg",
     "Figure 2. The inventory of record: fifteen stacks, all complete. The two newest are the "
     "delegated human seats - steward and deployer.")

h2("Who can block what")
para("The permission question (“who can do what”) has a twin that "
     "governance reviews usually miss: whose absence stops the system. The "
     "design answer is that almost nobody's should.")
diagram("criticality", "Figure 3. Only the consultant is critical to a run - their own checkpoints. "
                       "Only the FDE and platform engineer are critical to engineering. The steward "
                       "governs without blocking; the admin is dormant.")
table([
    ["Seat", "Surface", "Critical to"],
    ["Consultant", "Hosted console (Cognito)", "Execution - own checkpoints"],
    ["Forward-deployed engineer", "Workbench terminal + Claude Code", "Engineering"],
    ["Platform engineer", "Deployer seat (CloudFormation only)", "Engineering"],
    ["Data steward", "Steward seat (admission + audit)", "Nothing at runtime"],
    ["Analyst", "Databricks, read-only", "Nothing"],
    ["Account admin", "SSO, dormant", "Nothing, by design"],
], widths=[1.9, 2.6, 1.9])
doc.add_page_break()

# --- consultant ---------------------------------------------------------------
h1("1. Consultant / Engagement Lead")
para("Surface: the hosted console in a browser. No AWS account and no cloud "
     "permissions - a Cognito login created by the data steward. The consultant "
     "is the only human a running report ever waits for.", bold=True)

h2("Activity 1.1 - Run a report")
steps([
    "Sign in at the console URL. There is no self-signup; ask a steward for a login.",
    "Pick a client from the list, or type any company under Research Any Company.",
    "Entity verification runs first - a nonsense company is rejected for a few hundred tokens instead of tens of thousands.",
    "Watch progress live, or close the browser: run state is durable and consumes nothing while paused.",
])
shot("15-consultant-console.jpg",
     "Figure 4. The consultant home: twenty synthetic clients with workforce data, plus free-form "
     "company research.")

h2("Activity 1.2 - Decide at a checkpoint")
diagram("run", "Figure 5. A report run end to end. The two checkpoint pauses use waitForTaskToken, "
               "so a run waiting on a human costs nothing for hours or days.")
steps([
    "At the midpoint checkpoint, read the draft and choose: approve, request a revision, or reject.",
    "A revision takes your written feedback and injects it into the regeneration prompts - be concrete (“the attrition section contradicts the composition chart” beats “make it better”).",
    "Revisions are bounded (default two per checkpoint), so a run cannot loop forever.",
    "At the final checkpoint, approve to publish the artifacts.",
])
shot("12-execution-graph.png",
     "Figure 6. A real execution: stages in green flowing into MidpointApproval and FinalApproval.",
     width=4.6)

h2("What a consultant never has to worry about")
bullets([
    "Their login grants no cloud permissions; the app acts for them under its own scoped identity.",
    "They cannot break the pipeline, the data tiers, or anyone else's run from the console.",
    "Coarse-looking breakdowns are the privacy floor working - cells below five people are suppressed upstream.",
])
doc.add_page_break()

# --- FDE ----------------------------------------------------------------------
h1("2. Forward-Deployed Engineer (FDE)")
para("Surface: the workbench - an EC2 instance inside the account carrying Claude "
     "Code, git, the AWS and Databricks CLIs and Terraform. Reached over SSM "
     "only; no SSH key exists anywhere in the estate.", bold=True)

h2("Activity 2.1 - Reach the box")
para("Two identities are in play, and keeping them apart is the whole design. "
     "The FDE seat (intelligence-engine-dev-fde) is the human's right to reach "
     "the workbench and nothing else: find it, wake it, open a session, sleep it. "
     "Everything the engineer can actually do arrives from the instance role once "
     "they are at the prompt. Widening an FDE's power is therefore a reviewed "
     "template change, never a grant handed to a person.")
table([
    ["Path to the prompt", "What it needs", "When to use it"],
    ["Local terminal", "AWS CLI + Session Manager plugin (one-time install)", "Day to day - best ergonomics"],
    ["Console → Session Manager", "Console access + the FDE seat. No install", "A borrowed machine; no local setup"],
    ["CloudShell", "Also cloudshell:* - which the FDE seat does not grant", "Only if your permission set has it"],
], widths=[1.7, 2.6, 2.1])
code("aws ec2 start-instances --instance-ids <workbench-id> --profile intelligence-fde\n"
     "aws ssm start-session   --target <workbench-id>       --profile intelligence-fde\n"
     "sudo su - ec2-user      # lands in ~/intelligence_architecture, venv active\n"
     "claude")
para("Claude Code on the workbench runs on Amazon Bedrock through the instance "
     "role - no API key, no personal login. Agent inference therefore stays "
     "inside the account boundary: same Bedrock account, same CloudTrail, same "
     "data path as the product's own model calls.", italic=True)

h2("Activity 2.2 - Extend the behavioural eval (a full loop)")
para("The golden replay tests code. It says nothing about whether a *model* still "
     "behaves. The seeded-defect eval closes that gap: planted defects, one per "
     "failure class the advisory reviewer exists to catch, scored as recall. "
     "Extending it is ordinary FDE work, and the loop below was executed on the "
     "live workbench.")
steps([
    "Ask Claude Code for the change: add a fourth defect class - a year-over-year trend asserted from point-in-time counts.",
    "Read the diff it produced, in the terminal, before running anything.",
    "Run the eval against live Bedrock - four real model calls.",
    "Read the recall line. That number is evidence about the model, not the code.",
    "Keep or revert. Below the 2/3 floor is a regression even when every unit test passes.",
])
shot("18-fde-diff-stale-case.jpg",
     "Figure 7. Step 2: the agent's diff. A new 'stale' case with counts, a summary asserting a "
     "trend, and the guard test updated from three cases to four.")
shot("19-fde-eval-recall.jpg",
     "Figure 8. Steps 3-4: the eval against live Bedrock. 'reviewer recall 100% - caught "
     "[arithmetic, unsupported, contradiction, stale], missed []'. Note the caching warning above "
     "it - the wrapper reports honestly that the prefix is too short to cache.")

h2("Activity 2.3 - Investigate model cost, and probe the boundary")
para("Two more activities from the same session, both in the familiar Claude Code "
     "interface. First a cost question raised by that caching warning; then a "
     "direct test of what this seat may and may not do.")
steps([
    "Ask why the caching warning fired and whether caching is worth pursuing for this workload.",
    "Ask the agent to probe its own boundary: list the derived tier, then try to update a CloudFormation stack.",
    "Approve each shell command at the permission prompt - the agent cannot run commands unattended by default.",
    "Watch it self-correct: the first attempts used a named profile that does not exist on the box, so it checked its caller identity and retried against the instance role.",
    "Read the verdict table it produces.",
])
shot("20-fde-claude-code-session.jpg",
     "Figure 9. The whole session. Top: why prompt caching does not apply below 4,096 tokens and "
     "when it would be worth it. Bottom: the boundary probe - reading derived/ succeeded, "
     "cloudformation:UpdateStack was denied - with the agent's own explanation of why that split "
     "is intentional.")
para("That denial is the access model asserting itself against the engineer's own "
     "agent. The workbench changes code; only the deployer seat changes "
     "infrastructure.", italic=True)

h2("Activity 2.4 - Ship a pipeline version")
diagram("release", "Figure 10. Candidate to blessed. Every run pins an image digest, so which "
                   "pipeline version produced a report is never ambiguous.")
code("./infrastructure/remote_build.sh --stage-image        # builds in CodeBuild; prints digest\n"
     "python scripts/start_harness_run.py --client sterling-pharma \\\n"
     "       --stage-image <repo-uri>@sha256:<candidate>    # blessed digest still serves everyone else")
para("When the candidate passes, the digest is handed to the deployer seat for "
     "promotion. Releasing is deliberately not an FDE action - the same "
     "separation as infrastructure, applied to the pipeline default.")

h2("Activity 2.5 - Know where the work shows up")
para("The terminal is one surface, not the whole job.")
table([
    ["Question", "Surface that answers it"],
    ["What did I change?", "git branch and diff → the pull request"],
    ["Which pipeline version ran?", "Step Functions executions - each input pins an image digest"],
    ["What did a stage actually do?", "CodeBuild build history and its log stream"],
    ["What did the model see?", "/engineer run traces - the exact injected context"],
    ["What did it produce?", "analysis/<client_id>/ in the artifacts bucket"],
    ["What did it cost in tokens?", "CloudWatch Bedrock dashboard; summary() per run"],
    ["Is the estate still as described?", "python scripts/iac_coverage.py"],
    ["Who did what, account-wide?", "CloudTrail - including the agent's own calls"],
], widths=[2.6, 3.8])
shot("16-engineer-dashboard.jpg",
     "Figure 11. The engineer console: run states, checkpoints, datasets, infrastructure. An "
     "observation deck, not a cockpit.")
shot("17-guardrails.jpg",
     "Figure 12. The guardrail surface - input validation, entity verification, output validation, "
     "cost ceilings, data-access isolation. Read-only when deployed: enforcement changes ride "
     "through git.")
doc.add_page_break()

# --- platform engineer --------------------------------------------------------
h1("3. Platform Engineer")
para("Surface: the deployer seat. Operates CloudFormation on project stacks "
     "through cfn-exec - an execution role whose trust policy admits only the "
     "CloudFormation service, never a human - and mutates nothing directly.", bold=True)
diagram("deploy", "Figure 13. The deploy channel. Boundaries change through exactly one path: a "
                  "reviewed template, applied as a logged stack operation.")

h2("Activity 3.1 - Routine deploy")
code("./infrastructure/deploy.sh --profile intelligence-deployer \\\n"
     "  --cfn-role arn:aws:iam::<account>:role/intelligence-engine-dev-cfn-exec")
para("Quirks worth knowing before they surprise you: the app stack deploys twice "
     "(it must learn its own URL before Cognito can trust it); the steward stack "
     "takes the resulting user-pool id; unspecified parameters reuse their "
     "previous values on update.")

h2("Activity 3.2 - Verify, every time")
code("python scripts/qa_sweep.py     --profile intelligence-deployer\n"
     "python scripts/iac_coverage.py --profile intelligence-deployer")
para("Both are read-only and exit non-zero on failure, so either can gate a "
     "pipeline. An unmanaged resource is an incident, not housekeeping: adopt it "
     "into a template or remove it, and find out which seat created it.")

h2("Activity 3.3 - Promote or roll back a pipeline version")
code("python scripts/promote_stage_image.py --profile intelligence-deployer --digest sha256:...")
para("The script refuses a digest that is not in ECR, records the blessed image "
     "in SSM, and prints the previous digest - rollback is running it again with "
     "that value. Promotion affects the next run, never a run already in flight.")

h2("Activity 3.4 - Roll back infrastructure")
code("git checkout <last-good-commit> -- infrastructure/cloudformation/<name>.yaml\n"
     "aws cloudformation deploy ...    # the template is the truth you want back")
para("A stack stuck in ROLLBACK_COMPLETE after a failed create must be deleted "
     "and recreated - CloudFormation will not update it.")
doc.add_page_break()

# --- steward ------------------------------------------------------------------
h1("4. Data Steward")
para("Surface: the steward seat. Exactly three capabilities - admit data, admit "
     "people, read the audit surface - and no ability to deploy, write a governed "
     "tier, invoke a model, or reach the vault.", bold=True)
diagram("admission", "Figure 14. Admission through the medallion tiers. The steward decides what "
                     "enters; the gates in code enforce the rules on every write thereafter.")

h2("Activity 4.1 - Admit a source")
code("aws s3 cp ./drop.jsonl s3://<landing-bucket>/landing/<client>/<source>/<date>/ \\\n"
     "  --profile intelligence-steward")
para("The steward seat is the only identity that can write landing/, so every "
     "object there is something a steward deliberately placed. Licensing, privacy "
     "basis and inferred-attribute sensitivity are decisions taken before this "
     "command, not after.")
shot("02-lakehouse-tiers.jpg",
     "Figure 15. The governed data plane: medallion tiers as prefixes - foundational "
     "(pseudonymous), derived (aggregates, small cells suppressed), contextualized (vectors and "
     "graph), stewardship (append-only log).")

h2("Activity 4.2 - Admit and remove people")
code("aws cognito-idp admin-create-user       --profile intelligence-steward ...\n"
     "aws cognito-idp admin-set-user-password --profile intelligence-steward ... --permanent\n"
     "aws cognito-idp admin-disable-user      --profile intelligence-steward ...")
para("Set a permanent password rather than relying on the temporary-password "
     "flow: no email delivery is configured, so an invited user would otherwise "
     "be stranded.")

h2("Activity 4.3 - Read the audit surface")
diagram("audit", "Figure 16. Append-only by construction: the identities that write the log cannot "
                 "read it, and the identity that reads it can write nothing but the drop zone.")
shot("05-stewardship-gate-events.jpg",
     "Figure 17. The stewardship log in S3. gate-events/ records every gate evaluation as JSONL - "
     "which gate, which caller, BLOCK / WARN / NOTE, and what triggered it.")
shot("13-cloudtrail-trail.jpg",
     "Figure 18. CloudTrail: multi-region, log-file validation enabled, and object-level data "
     "events on the governed prefixes - every individual read of governed data, not just "
     "management calls.")

h2("Activity 4.4 - Review a lineage claim")
shot("07-glue-lineage-properties.jpg",
     "Figure 19. Governance as table metadata: ie.owner names the accountable human, lineage "
     "records 500 rows in and 37 aggregate cells out, and contains_personal_data=false sits beside "
     "derived_from_personal_data=true - the honest pair.")
doc.add_page_break()

# --- analyst ------------------------------------------------------------------
h1("5. Data Scientist / Analyst")
para("Surface: Databricks, reading the lakehouse in place through Unity Catalog. "
     "Zero-copy, read-only, and no AWS credentials in the analyst's hands.", bold=True)
diagram("analyst", "Figure 20. The zero-copy path. No write path exists from the notebook back to "
                   "the governed tiers.")
bullets([
    "derived/ is the everyday surface for BI - a normal external table with suppression already applied.",
    "foundational/ supports record-grain modelling, and carries record-grain obligations: pseudonymous is not anonymous.",
    "contextualized/ ships embeddings beside their metadata columns for retrieval prototyping.",
    "Reconstructing suppressed cells by joining or differencing derived outputs defeats the control the tier exists to provide.",
])
para("New or changed product tables are requested, not built here: they are a "
     "reviewed change to the build stages, run under the product-builder identity "
     "with lineage metadata updated. That is what keeps the catalog's ownership "
     "claims true.")
doc.add_page_break()

# --- admin --------------------------------------------------------------------
h1("6. Account Admin")
para("Dormant by design. Two jobs only: seating people, and recovery if the "
     "deploy channel itself breaks. Used on a normal day means something is "
     "misdesigned.", bold=True)
h2("Activity 6.1 - Seat a person")
para("Each seat is a role trusting the account; seating is one statement granting "
     "a principal permission to assume it. Unseating is removing that statement.")
code('{"Effect": "Allow", "Action": "sts:AssumeRole",\n'
     ' "Resource": "arn:aws:iam::<account>:role/intelligence-engine-dev-steward"}')
para("The FDE is seated slightly differently - through the fde role, which itself "
     "carries the SSM and EC2 permissions needed to reach the workbench. "
     "Consultants are not seated by the admin at all; they are Cognito users the "
     "steward manages.")
h2("Activity 6.2 - Break glass, then reconcile")
para("Fix the minimum, through a template wherever possible, then prove the "
     "estate still matches the repository. Break-glass that leaves untracked "
     "state behind is not recovery, it is new damage.")
code("python scripts/iac_coverage.py   # must report NO DRIFT afterwards")
doc.add_page_break()

# --- appendix -----------------------------------------------------------------
h1("Appendix A - The harness in detail")
shot("09-stepfunctions-machine.jpg",
     "Figure 21. The report-build state machine and its execution history.")
shot("10-stepfunctions-definition.jpg",
     "Figure 22. Definition beside graph. Execution input pins stage_image by digest; approvals "
     "carry {decision, feedback}; revision loops are bounded by max_revisions.")

h1("Appendix B - The seats, as IAM sees them")
table([
    ["Role", "Held by", "Can", "Cannot"],
    ["fde", "A human, via AssumeRole", "Find, wake, reach, sleep the workbench", "Read tiers, invoke models, deploy"],
    ["workbench", "The EC2 instance", "Bedrock, tier reads, artifacts, builds, harness runs", "Deploy, touch IAM or billing"],
    ["steward", "A human, via AssumeRole", "landing-write, Cognito users, read the audit log", "Write governed tiers, deploy, vault"],
    ["deployer", "A human, via AssumeRole", "CloudFormation on project stacks via cfn-exec", "Mutate any resource directly"],
    ["cfn-exec", "CloudFormation only", "Create and update stack resources", "Be assumed by any human"],
    ["conformance", "CodeBuild", "landing-read → foundational-write", "Read derived or contextualized"],
    ["product-builder", "CodeBuild", "foundational-read → derived + contextualized write", "Write foundational"],
    ["stage-runner", "CodeBuild", "Tier reads, artifact write, Bedrock", "Write any governed tier"],
    ["databricks-uc", "Databricks", "Read the lakehouse in place", "Write anything"],
], widths=[1.3, 1.5, 2.1, 1.7])

h1("Appendix C - Verifying any claim in this document")
para("Nothing here needs to be taken on trust; each claim has a command.")
code("python scripts/qa_sweep.py       # stacks, auth, tiers, harness, idle billing\n"
     "python scripts/iac_coverage.py   # nothing live outside the templates\n"
     "python -m pytest                 # the code suite\n"
     "python -m pytest tests/test_seeded_defects.py -s   # model behaviour, live Bedrock\n"
     "aws iam list-attached-role-policies --role-name intelligence-engine-dev-stage-runner")

h1("Appendix D - Where everything lives")
table([
    ["Artefact", "Location"],
    ["Repository", "github.com/dorchester/intelligence_architecture"],
    ["Role guides (full command detail)", "docs/guides/"],
    ["Access model, every identity and grant", "docs/access-model.md"],
    ["Architecture diagrams (source of Figures 1, 3, 5, 10, 13, 14, 16, 20)", "docs/diagrams.md"],
    ["Deploy-from-nothing walkthrough", "docs/aws-walkthrough.md"],
    ["Substrate / workload boundary contract", "docs/integration-contract.md"],
    ["This document's generator", "scripts/build_role_guide_docx.py"],
    ["Diagram renderer", "scripts/render_diagrams.py"],
    ["Screenshot trim / redact utility", "scripts/prepare_screenshots.py"],
], widths=[3.4, 3.0])

doc.save(args.out)
print("saved:", args.out)
